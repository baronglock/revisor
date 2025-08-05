from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table
import re
from copy import deepcopy

class WordDocumentHandler:
    """Classe para manipulação de documentos Word preservando formatação"""
    
    @staticmethod
    def read_document_complete(doc_path):
        """Lê documento preservando TODA estrutura, formatação e conteúdo"""
        doc = Document(doc_path)
        content = []
        
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Parágrafo
                para = Paragraph(element, doc)
                content.append({
                    'type': 'paragraph',
                    'text': para.text,
                    'style': para.style.name if para.style else None,
                    'alignment': para.alignment,
                    'runs': WordDocumentHandler._extract_runs(para),
                    'element': element
                })
            elif element.tag.endswith('tbl'):  # Tabela
                table = Table(element, doc)
                content.append({
                    'type': 'table',
                    'data': WordDocumentHandler._extract_table_data(table),
                    'element': element
                })
        
        return content
    
    @staticmethod
    def _extract_runs(paragraph):
        """Extrai runs com formatação detalhada"""
        runs_data = []
        for run in paragraph.runs:
            run_data = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size.pt if run.font.size else None,
                'color': None
            }
            
            # Verifica se é um link
            if WordDocumentHandler._is_hyperlink(run):
                run_data['is_link'] = True
            
            # Cor do texto
            if run.font.color and run.font.color.rgb:
                run_data['color'] = str(run.font.color.rgb)
            
            runs_data.append(run_data)
        
        return runs_data
    
    @staticmethod
    def _is_hyperlink(run):
        """Verifica se o run contém um hyperlink"""
        # Verificação simplificada por padrão de URL
        url_pattern = r'https?://[^\s]+'
        return bool(re.search(url_pattern, run.text))
    
    @staticmethod
    def _extract_table_data(table):
        """Extrai dados da tabela preservando formatação"""
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_content = []
                for paragraph in cell.paragraphs:
                    cell_content.append({
                        'text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else None,
                        'runs': WordDocumentHandler._extract_runs(paragraph)
                    })
                row_data.append(cell_content)
            table_data.append(row_data)
        return table_data
    
    @staticmethod
    def create_document_from_content(content, revised_texts=None):
        """Recria documento preservando estrutura e aplicando revisões"""
        new_doc = Document()
        
        # Remove o parágrafo padrão vazio
        if new_doc.paragraphs:
            p = new_doc.paragraphs[0]._element
            p.getparent().remove(p)
        
        text_index = 0
        
        for item in content:
            if item['type'] == 'paragraph':
                # Usa texto revisado se disponível
                if revised_texts and text_index < len(revised_texts):
                    new_text = revised_texts[text_index]
                    text_index += 1
                else:
                    new_text = item['text']
                
                # Cria parágrafo com estilo original
                para = new_doc.add_paragraph()
                if item['style']:
                    para.style = item['style']
                if item['alignment']:
                    para.alignment = item['alignment']
                
                # Aplica formatação dos runs
                WordDocumentHandler._apply_runs_to_paragraph(
                    para, item['runs'], new_text
                )
                
            elif item['type'] == 'table':
                # Recria tabela
                table_data = item['data']
                if table_data:
                    rows = len(table_data)
                    cols = len(table_data[0]) if table_data[0] else 0
                    table = new_doc.add_table(rows=rows, cols=cols)
                    
                    for r_idx, row_data in enumerate(table_data):
                        for c_idx, cell_data in enumerate(row_data):
                            cell = table.cell(r_idx, c_idx)
                            # Limpa célula
                            for p in cell.paragraphs:
                                p._element.getparent().remove(p._element)
                            
                            # Adiciona parágrafos na célula
                            for para_data in cell_data:
                                para = cell.add_paragraph()
                                if para_data.get('style'):
                                    para.style = para_data['style']
                                
                                # Usa texto revisado se disponível
                                if revised_texts and text_index < len(revised_texts):
                                    new_text = revised_texts[text_index]
                                    text_index += 1
                                else:
                                    new_text = para_data['text']
                                
                                WordDocumentHandler._apply_runs_to_paragraph(
                                    para, para_data['runs'], new_text
                                )
        
        return new_doc
    
    @staticmethod
    def _apply_runs_to_paragraph(paragraph, original_runs, new_text):
        """Aplica formatação dos runs ao novo texto"""
        # Se não há runs ou texto vazio
        if not original_runs or not new_text:
            paragraph.add_run(new_text)
            return
        
        # Detecta e preserva links
        url_pattern = r'(https?://[^\s]+)'
        parts = re.split(url_pattern, new_text)
        
        current_pos = 0
        for part in parts:
            if not part:
                continue
                
            run = paragraph.add_run(part)
            
            # Se é URL, aplica formatação de link
            if re.match(url_pattern, part):
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
            else:
                # Aplica formatação do run original mais próximo
                if original_runs:
                    # Usa formatação do primeiro run como base
                    orig_run = original_runs[0]
                    if orig_run.get('bold'):
                        run.bold = True
                    if orig_run.get('italic'):
                        run.italic = True
                    if orig_run.get('underline') and not re.match(url_pattern, part):
                        run.underline = True
                    if orig_run.get('font_name'):
                        run.font.name = orig_run['font_name']
                    if orig_run.get('font_size'):
                        run.font.size = Pt(orig_run['font_size'])
                    if orig_run.get('color') and not re.match(url_pattern, part):
                        try:
                            color_hex = orig_run['color']
                            if isinstance(color_hex, str) and len(color_hex) == 6:
                                run.font.color.rgb = RGBColor(
                                    int(color_hex[0:2], 16),
                                    int(color_hex[2:4], 16),
                                    int(color_hex[4:6], 16)
                                )
                        except:
                            pass
    
    @staticmethod
    def extract_images_info(doc_path):
        """Extrai informações sobre imagens no documento"""
        doc = Document(doc_path)
        images_info = []
        
        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.target_ref:
                images_info.append({
                    'index': i,
                    'relationship_id': rel.rId,
                    'target': rel.target_ref
                })
        
        return images_info
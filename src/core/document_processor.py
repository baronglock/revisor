import os
import shutil
import logging
import json
import re
from typing import List, Dict, Tuple
from docx import Document
from ..utils.word_utils import WordDocumentHandler
from ..utils.api_client import OpenAIClient

class DocumentProcessor:
    """Processa documentos Word identificando e corrigindo apenas erros"""
    
    def __init__(self, api_key: str, model: str = "gpt-4.1"):
        self.api_client = OpenAIClient(api_key, model)
        self.api_key = api_key
        self.model = model
        self.word_handler = WordDocumentHandler()
        self.logger = logging.getLogger(__name__)
        
        # Define tamanho de chunk baseado no modelo
        if model == "gpt-4.1":
            self.max_chunk_size = 10000  # 800k chars para GPT-4.1 (processa tudo de uma vez)
        else:
            self.max_chunk_size = 10000  # 200k chars para outros modelos
    
    
    
    def process_document(self, input_path: str, output_path: str, callback=None):
        """Processa documento com precisão MÁXIMA"""
        try:
            # 1. Copia o arquivo original
            self.logger.info(f"Iniciando processamento ULTRA-PRECISO")
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            shutil.copy2(input_path, output_path)
            
            # 2. Abre AMBOS os documentos
            original_doc = Document(input_path)
            doc = Document(output_path)
            
            # 3. Mapeia TODOS os textos com índices CORRETOS
            all_paragraphs = []
            paragraph_counter = 0
            
            # Parágrafos normais
            for i, (orig_para, para) in enumerate(zip(original_doc.paragraphs, doc.paragraphs)):
                if para.text.strip():
                    paragraph_counter += 1
                    all_paragraphs.append({
                        'global_index': len(all_paragraphs),
                        'paragraph_number': paragraph_counter,
                        'doc_index': i,
                        'original_text': orig_para.text,
                        'current_text': para.text,
                        'paragraph_obj': para,
                        'type': 'normal',
                        'location': f'Parágrafo {paragraph_counter}',
                        'page_estimate': paragraph_counter // 3  # ~3 parágrafos por página
                    })
            
            # Tabelas
            for t_idx, (orig_table, table) in enumerate(zip(original_doc.tables, doc.tables)):
                for r_idx, (orig_row, row) in enumerate(zip(orig_table.rows, table.rows)):
                    for c_idx, (orig_cell, cell) in enumerate(zip(orig_row.cells, row.cells)):
                        for p_idx, (orig_para, para) in enumerate(zip(orig_cell.paragraphs, cell.paragraphs)):
                            if para.text.strip():
                                paragraph_counter += 1
                                key = f"table_{t_idx}_{r_idx}_{c_idx}_{p_idx}"
                                all_paragraphs.append({
                                    'global_index': len(all_paragraphs),
                                    'paragraph_number': paragraph_counter,
                                    'doc_index': key,
                                    'original_text': orig_para.text,
                                    'current_text': para.text,
                                    'paragraph_obj': para,
                                    'type': 'table',
                                    'location': f'Tabela {t_idx+1}, Célula ({r_idx+1},{c_idx+1})',
                                    'page_estimate': paragraph_counter // 3
                                })
            
            self.logger.info(f"Total de {len(all_paragraphs)} parágrafos para análise DETALHADA")
            
            # 4. Cria blocos PEQUENOS para máxima precisão
            blocks = self._create_precise_blocks(all_paragraphs)
            self.logger.info(f"Dividido em {len(blocks)} blocos pequenos para análise minuciosa")
            
            # 5. Processa CADA bloco com atenção total
            all_corrections = []
            total_corrections_applied = 0
            
            for block_idx, block in enumerate(blocks):
                # Informação clara sobre o bloco
                first_para = block[0]['paragraph_number']
                last_para = block[-1]['paragraph_number']
                page_range = f"páginas {block[0]['page_estimate']+1}-{block[-1]['page_estimate']+1}"
                
                if callback:
                    callback(block_idx + 1, len(blocks), 
                            f"Analisando parágrafos {first_para}-{last_para} ({page_range})")
                
                self.logger.info(f"Bloco {block_idx+1}/{len(blocks)}: "
                            f"Parágrafos {first_para}-{last_para} ({len(block)} textos)")
                
                # Prepara texto para análise MINUCIOSA
                block_text = self._prepare_block_for_analysis(block)
                
                # Envia para análise
                corrections = self.api_client.identify_errors_precise(block_text, block_idx)
                
                if corrections:
                    self.logger.info(f"Bloco {block_idx+1}: {len(corrections)} erros encontrados")
                    
                    # Aplica CADA correção
                    for corr in corrections:
                        # Encontra o parágrafo correto
                        para_data = self._find_paragraph_in_block(block, corr)
                        
                        if para_data:
                            # Aplica a correção
                            success = self._apply_correction_ultra_precise(para_data, corr)
                            
                            if success:
                                total_corrections_applied += 1
                                
                                # Registra correção completa
                                all_corrections.append({
                                    'block': block_idx + 1,
                                    'paragraph_number': para_data['paragraph_number'],
                                    'location': para_data['location'],
                                    'page': para_data['page_estimate'] + 1,
                                    'error': corr.get('error', ''),
                                    'correction': corr.get('correction', ''),
                                    'type': corr.get('type', 'outros'),
                                    'original_text': para_data['original_text'],
                                    'corrected_text': para_data['paragraph_obj'].text,
                                    'applied': True
                                })
                            else:
                                self.logger.warning(f"Falha ao aplicar: {corr}")
            
            # 6. Verifica se TODAS as mudanças foram detectadas
            self.logger.info("Verificação final de integridade...")
            
            for para_data in all_paragraphs:
                current = para_data['paragraph_obj'].text
                original = para_data['original_text']
                
                if current != original:
                    # Verifica se foi registrado
                    found = any(c['paragraph_number'] == para_data['paragraph_number'] 
                            for c in all_corrections)
                    
                    if not found:
                        # Mudança não detectada!
                        self.logger.warning(f"Mudança não registrada no parágrafo {para_data['paragraph_number']}")
                        
                        diff = self._analyze_difference(original, current)
                        all_corrections.append({
                            'block': 'auto',
                            'paragraph_number': para_data['paragraph_number'],
                            'location': para_data['location'],
                            'page': para_data['page_estimate'] + 1,
                            'error': diff['error'],
                            'correction': diff['correction'],
                            'type': 'auto-detectado',
                            'original_text': original,
                            'corrected_text': current,
                            'applied': True
                        })
            
            # 7. Salva documento
            doc.save(output_path)
            self.logger.info(f"Documento salvo com {len(all_corrections)} correções totais")
            
            # 8. Salva relatório detalhado
            report_path = self._save_detailed_report(output_path, all_corrections)
            
            return output_path
            
        except Exception as e:
            self.logger.error(f"Erro: {str(e)}")
            raise


    def _create_precise_blocks(self, all_paragraphs: List[Dict]) -> List[List[Dict]]:
        """Cria blocos PEQUENOS para análise precisa"""
        blocks = []
        current_block = []
        current_size = 0
        
        for para_data in all_paragraphs:
            text_size = len(para_data['current_text'])
            
            # Força novo bloco se muito grande ou muitos parágrafos
            if (current_size + text_size > self.max_chunk_size or 
                len(current_block) >= 50):  # Máximo 50 parágrafos por bloco
                
                if current_block:
                    blocks.append(current_block)
                current_block = [para_data]
                current_size = text_size
            else:
                current_block.append(para_data)
                current_size += text_size
        
        if current_block:
            blocks.append(current_block)
        
        return blocks

    def _prepare_block_for_analysis(self, block: List[Dict]) -> str:
        """Prepara bloco com contexto MÁXIMO para análise"""
        block_text = f"BLOCO DE PARÁGRAFOS {block[0]['paragraph_number']} a {block[-1]['paragraph_number']}:\n\n"
        
        for para_data in block:
            # Adiciona contexto completo
            block_text += f"[PARÁGRAFO {para_data['paragraph_number']}]\n"
            block_text += f"[LOCALIZAÇÃO: {para_data['location']}]\n"
            
            # Marca tipo de conteúdo
            text = para_data['current_text']
            if len(text) < 100 and not text.endswith(('.', '!', '?', ':')):
                block_text += "[TIPO: TÍTULO/CABEÇALHO]\n"
            elif text.strip().startswith(('•', '-', '1.', '2.', 'a)', 'b)')):
                block_text += "[TIPO: ITEM DE LISTA]\n"
            elif para_data['type'] == 'table':
                block_text += "[TIPO: CÉLULA DE TABELA]\n"
            else:
                block_text += "[TIPO: PARÁGRAFO NORMAL]\n"
            
            block_text += f"{text}\n"
            block_text += f"[FIM_PARÁGRAFO_{para_data['paragraph_number']}]\n\n"
        
        return block_text

    def _find_paragraph_in_block(self, block: List[Dict], correction: Dict) -> Dict:
        """Encontra parágrafo exato da correção"""
        error_text = correction.get('error', '')
        para_num = correction.get('paragraph', 0)
        
        # Tenta pelo número do parágrafo primeiro
        if para_num > 0:
            for para_data in block:
                if para_data['paragraph_number'] == para_num:
                    return para_data
        
        # Tenta pelo conteúdo do erro
        for para_data in block:
            if error_text in para_data['current_text']:
                return para_data
        
        return None

    def _apply_correction_ultra_precise(self, para_data: Dict, correction: Dict) -> bool:
        """Aplica correção com precisão máxima"""
        try:
            paragraph = para_data['paragraph_obj']
            original_text = paragraph.text
            error = correction.get('error', '')
            fix = correction.get('correction', '')
            
            if not error or not fix:
                return False
            
            # Log detalhado
            self.logger.debug(f"Aplicando no parágrafo {para_data['paragraph_number']}: '{error}' → '{fix}'")
            
            # Múltiplas estratégias de substituição
            new_text = original_text
            
            # 1. Substituição exata
            if error in original_text:
                new_text = original_text.replace(error, fix, 1)
            
            # 2. Com word boundaries
            elif re.search(r'\b' + re.escape(error) + r'\b', original_text):
                new_text = re.sub(r'\b' + re.escape(error) + r'\b', fix, original_text, count=1)
            
            # 3. Ignorando case
            elif error.lower() in original_text.lower():
                # Substituição case-insensitive preservando o case original quando possível
                import re
                pattern = re.compile(re.escape(error), re.IGNORECASE)
                new_text = pattern.sub(fix, original_text, count=1)
            
            # 4. Tentativa com espaços flexíveis
            else:
                # Remove espaços extras e tenta novamente
                normalized_error = ' '.join(error.split())
                normalized_text = ' '.join(original_text.split())
                if normalized_error in normalized_text:
                    # Encontra e substitui considerando espaços flexíveis
                    new_text = original_text.replace(error.strip(), fix.strip(), 1)
            
            # Verifica se conseguiu aplicar
            if new_text != original_text:
                paragraph.text = new_text
                self.logger.info(f"✓ Correção aplicada no parágrafo {para_data['paragraph_number']}")
                return True
            else:
                self.logger.warning(f"✗ Não conseguiu aplicar no parágrafo {para_data['paragraph_number']}: '{error}'")
                return False
                
        except Exception as e:
            self.logger.error(f"Erro ao aplicar correção: {str(e)}")
            return False

    def _analyze_difference(self, original: str, current: str) -> Dict:
        """Analisa diferença entre textos com precisão"""
        # Casos especiais comuns
        if current == original + '.':
            return {'error': '[faltava ponto final]', 'correction': '.', 'type': 'pontuação'}
        
        if current == original + ',':
            return {'error': '[faltava vírgula]', 'correction': ',', 'type': 'pontuação'}
        
        if original.lower() == current.lower() and original[0].lower() == current[0].upper():
            return {'error': original[0], 'correction': current[0], 'type': 'maiúscula'}
        
        # Análise detalhada
        import difflib
        s = difflib.SequenceMatcher(None, original.split(), current.split())
        
        diffs = []
        for tag, i1, i2, j1, j2 in s.get_opcodes():
            if tag != 'equal':
                error_text = ' '.join(original.split()[i1:i2]) if i1 < i2 else '[vazio]'
                fix_text = ' '.join(current.split()[j1:j2]) if j1 < j2 else '[removido]'
                diffs.append({'error': error_text, 'correction': fix_text})
        
        if diffs:
            return diffs[0]
        
        return {'error': 'diferença sutil', 'correction': 'texto alterado', 'type': 'outros'}

    def _find_difference(self, original: str, corrected: str) -> Dict:
        """Encontra a diferença entre dois textos"""
        import difflib
        
        # Se a diferença é só pontuação no final
        if corrected == original + '.':
            return {
                'error': '[falta ponto final]',
                'correction': '.',
                'type': 'pontuação'
            }
        
        # Usa SequenceMatcher para encontrar diferenças
        s = difflib.SequenceMatcher(None, original, corrected)
        
        for tag, i1, i2, j1, j2 in s.get_opcodes():
            if tag == 'replace':
                return {
                    'error': original[i1:i2],
                    'correction': corrected[j1:j2],
                    'type': 'correção automática'
                }
            elif tag == 'insert':
                return {
                    'error': '[faltando]',
                    'correction': corrected[j1:j2],
                    'type': 'inserção'
                }
            elif tag == 'delete':
                return {
                    'error': original[i1:i2],
                    'correction': '[removido]',
                    'type': 'remoção'
                }
        
        return {
            'error': 'mudança não identificada',
            'correction': 'texto alterado',
            'type': 'outros'
        }

    def _save_complete_report(self, output_path: str, all_corrections: List[Dict], api_corrections: List[Dict]):
        """Salva relatório COMPLETO com todas as correções"""
        report_path = output_path.replace('.docx', '_complete_report.json')
        
        # Estatísticas detalhadas
        stats_by_type = {}
        stats_by_source = {'api': 0, 'auto_detected': 0}
        
        for corr in all_corrections:
            # Por tipo
            corr_type = corr.get('type', 'outros')
            stats_by_type[corr_type] = stats_by_type.get(corr_type, 0) + 1
            
            # Por fonte
            source = corr.get('source', 'unknown')
            stats_by_source[source] = stats_by_source.get(source, 0) + 1
        
        report = {
            'summary': {
                'total_corrections': len(all_corrections),
                'reported_by_api': len(api_corrections),
                'auto_detected': stats_by_source.get('auto_detected', 0),
                'by_type': stats_by_type,
                'by_source': stats_by_source
            },
            'all_corrections': all_corrections
        }
        
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        
        self.logger.info(f"Relatório completo salvo: {report_path}")
        return report_path
            
    

    
    def _extract_local_index(self, correction: Dict, block_text: str) -> int:
        """Extrai índice local do texto no bloco"""
        error_text = correction.get('error', '')
        
        # Procura o erro em cada texto marcado
        import re
        pattern = r'\[TEXTO_(\d+)\](.*?)\[FIM_TEXTO_\d+\]'
        matches = list(re.finditer(pattern, block_text, re.DOTALL))
        
        for match in matches:
            local_idx = int(match.group(1))
            text_content = match.group(2).strip()
            
            # Verifica se o erro está neste texto
            if error_text in text_content:
                return local_idx
        
        # Se não encontrou pelo erro exato, tenta pela linha indicada
        line_num = correction.get('line', 0)
        if line_num > 0:
            # Estima qual texto baseado na linha
            current_line = 0
            for match in matches:
                local_idx = int(match.group(1))
                text_lines = match.group(2).strip().count('\n') + 1
                if current_line <= line_num <= current_line + text_lines:
                    return local_idx
                current_line += text_lines
        
        return None
    
    def _apply_correction_with_verification(self, paragraph, correction: Dict) -> Tuple[bool, str, str]:
        """Aplica correção com verificação e retorna (sucesso, texto_original, texto_corrigido)"""
        try:
            original_text = paragraph.text
            error = correction.get('error', '')
            fix = correction.get('correction', '')
            
            if not error or not fix:
                return False, original_text, original_text
            
            # Tenta diferentes padrões de substituição
            new_text = original_text
            
            # 1. Tenta substituição exata com word boundaries
            pattern = r'\b' + re.escape(error) + r'\b'
            if re.search(pattern, original_text):
                new_text = re.sub(pattern, fix, original_text, count=1)
            
            # 2. Se não encontrou, tenta sem word boundaries (para casos com pontuação)
            elif error in original_text:
                new_text = original_text.replace(error, fix, 1)
            
            # 3. Se ainda não encontrou, tenta ignorar case
            else:
                pattern = r'\b' + re.escape(error) + r'\b'
                if re.search(pattern, original_text, re.IGNORECASE):
                    new_text = re.sub(pattern, fix, original_text, count=1, flags=re.IGNORECASE)
            
            # Verifica se houve mudança
            if new_text != original_text:
                paragraph.text = new_text
                self.logger.info(f"Correção aplicada: '{error}' → '{fix}'")
                return True, original_text, new_text
            else:
                self.logger.warning(f"Não foi possível aplicar: '{error}' → '{fix}' no texto: {original_text[:100]}...")
                return False, original_text, original_text
            
        except Exception as e:
            self.logger.error(f"Erro ao aplicar correção: {str(e)}")
            return False, paragraph.text, paragraph.text
    
    def _save_corrections_report(self, output_path: str, all_corrections: List[Dict], applied_corrections: List[Dict]):
        """Salva relatório detalhado das correções"""
        report_path = output_path.replace('.docx', '_corrections_report.json')
        
        # Estatísticas
        stats_by_type = {}
        for corr in applied_corrections:
            corr_type = corr.get('type', 'outros')
            stats_by_type[corr_type] = stats_by_type.get(corr_type, 0) + 1
        
        report = {
            'summary': {
                'total_errors_found': len(all_corrections),
                'total_corrections_applied': len(applied_corrections),
                'corrections_failed': len(all_corrections) - len(applied_corrections),
                'by_type': stats_by_type
            },
            'applied_corrections': applied_corrections,
            'all_corrections': all_corrections
        }
        
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        
        self.logger.info(f"Relatório detalhado salvo: {report_path}")
        return report_path
    
    

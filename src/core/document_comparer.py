import os
import json
import shutil
import logging
from typing import List, Dict
from docx import Document
from docx.shared import RGBColor
import difflib

class DocumentComparer:
    """Compara documentos preservando TODA formatação"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def compare_documents(self, original_path: str, revised_path: str, 
                         output_path: str, log_path: str = None) -> str:
        """Compara documentos usando comparação real de espelhamento"""
        
        # SEMPRE usa comparação completa espelhada
        return self.create_mirror_comparison(original_path, revised_path, output_path)
    
    def create_mirror_comparison(self, original_path: str, revised_path: str, output_path: str) -> str:
        """Cria comparação COPIANDO o documento revisado e marcando TODAS as diferenças"""
        try:
            self.logger.info("Criando comparação espelhada completa")
            
            # 1. COPIA o documento revisado (preserva TODA formatação)
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            shutil.copy2(revised_path, output_path)
            
            # 2. Abre os três documentos
            original_doc = Document(original_path)
            revised_doc = Document(revised_path)
            comparison_doc = Document(output_path)  # Abre a cópia
            
            # 3. Coleta TODAS as correções para o dicionário
            all_corrections = []
            para_num = 0
            
            # 4. Compara e marca diferenças em CADA parágrafo
            for orig_para, rev_para, comp_para in zip(original_doc.paragraphs, 
                                                      revised_doc.paragraphs, 
                                                      comparison_doc.paragraphs):
                
                if orig_para.text.strip():
                    para_num += 1
                    
                    if orig_para.text != rev_para.text:
                        # Houve mudança - analisa e registra
                        diff = self._analyze_paragraph_changes(orig_para.text, rev_para.text)
                        if diff:
                            all_corrections.append({
                                'paragraph_number': para_num,
                                'location': f'Parágrafo {para_num}',
                                'page': (para_num // 3) + 1,  # Estimativa ~3 parágrafos por página
                                'error': diff['error'],
                                'correction': diff['correction'],
                                'type': diff['type'],
                                'original_text': orig_para.text,
                                'corrected_text': rev_para.text
                            })
                        
                        # Marca no parágrafo do comparador
                        self._mark_paragraph_changes(comp_para, orig_para.text, rev_para.text)
            
            # 5. Compara e marca diferenças em TABELAS
            table_corrections = []
            for t_idx, (orig_table, rev_table, comp_table) in enumerate(zip(original_doc.tables, 
                                                                           revised_doc.tables, 
                                                                           comparison_doc.tables)):
                for r_idx, (orig_row, rev_row, comp_row) in enumerate(zip(orig_table.rows, 
                                                                         rev_table.rows, 
                                                                         comp_table.rows)):
                    for c_idx, (orig_cell, rev_cell, comp_cell) in enumerate(zip(orig_row.cells, 
                                                                                rev_row.cells, 
                                                                                comp_row.cells)):
                        for p_idx, (orig_para, rev_para, comp_para) in enumerate(zip(orig_cell.paragraphs, 
                                                                                    rev_cell.paragraphs, 
                                                                                    comp_cell.paragraphs)):
                            if orig_para.text != rev_para.text and orig_para.text.strip():
                                # Registra correção em tabela
                                diff = self._analyze_paragraph_changes(orig_para.text, rev_para.text)
                                if diff:
                                    table_corrections.append({
                                        'location': f'Tabela {t_idx+1}, Célula ({r_idx+1},{c_idx+1})',
                                        'page': (para_num // 3) + 1,
                                        'error': diff['error'],
                                        'correction': diff['correction'],
                                        'type': diff['type'],
                                        'original_text': orig_para.text,
                                        'corrected_text': rev_para.text
                                    })
                                
                                # Marca mudanças
                                self._mark_paragraph_changes(comp_para, orig_para.text, rev_para.text)
            
            # Adiciona correções de tabelas
            all_corrections.extend(table_corrections)
            
            # 6. Adiciona sumário e dicionário no INÍCIO do documento
            self._add_summary_with_dictionary(comparison_doc, all_corrections)
            
            # 7. Salva
            comparison_doc.save(output_path)
            self.logger.info(f"Comparação espelhada salva: {output_path}")
            self.logger.info(f"Total de {len(all_corrections)} correções documentadas")
            
            return output_path
            
        except Exception as e:
            self.logger.error(f"Erro ao criar comparação: {str(e)}")
            raise
    
    def _analyze_paragraph_changes(self, original: str, revised: str) -> Dict:
        """Analisa mudanças em um parágrafo e retorna diferença principal"""
        
        # Casos simples e comuns
        if revised == original + '.':
            return {'error': '[faltava ponto final]', 'correction': '.', 'type': 'pontuação'}
        
        if revised == original + ',':
            return {'error': '[faltava vírgula]', 'correction': ',', 'type': 'pontuação'}
        
        # Mudança de maiúscula no início
        if len(original) > 0 and len(revised) > 0:
            if original[0].islower() and revised[0].isupper() and original[1:] == revised[1:]:
                return {'error': original[0], 'correction': revised[0], 'type': 'maiúscula'}
        
        # Análise palavra por palavra
        orig_words = original.split()
        rev_words = revised.split()
        
        # Encontra primeira diferença
        for i, (o, r) in enumerate(zip(orig_words, rev_words)):
            if o != r:
                return {'error': o, 'correction': r, 'type': 'ortografia/gramática'}
        
        # Se adicionou palavras no final
        if len(rev_words) > len(orig_words):
            added = ' '.join(rev_words[len(orig_words):])
            return {'error': '[faltando]', 'correction': added, 'type': 'adição'}
        
        # Se removeu palavras
        if len(orig_words) > len(rev_words):
            removed = ' '.join(orig_words[len(rev_words):])
            return {'error': removed, 'correction': '[removido]', 'type': 'remoção'}
        
        # Mudança complexa
        return {'error': 'mudança complexa', 'correction': 'ver texto marcado', 'type': 'outros'}
    
    def _mark_paragraph_changes(self, paragraph, original_text: str, revised_text: str):
        """Marca mudanças em um parágrafo preservando formatação"""
        
        # Limpa o parágrafo mas preserva estilo
        para_style = paragraph.style
        para_alignment = paragraph.alignment
        
        # Remove runs existentes
        for run in paragraph.runs:
            run.text = ""
        
        # Casos especiais
        # Se só adicionou ponto final
        if revised_text == original_text + '.':
            paragraph.runs[0].text = original_text
            added = paragraph.add_run('.')
            added.font.color.rgb = RGBColor(0, 128, 0)
            added.bold = True
            return
        
        # Se mudou maiúscula no início
        if original_text.lower() == revised_text.lower() and original_text != revised_text:
            # Apenas mudança de case
            for i, (o_char, r_char) in enumerate(zip(original_text, revised_text)):
                if o_char != r_char:
                    # Marca o caractere mudado
                    if i > 0:
                        paragraph.add_run(revised_text[:i])
                    
                    changed = paragraph.add_run(r_char)
                    changed.font.color.rgb = RGBColor(0, 128, 0)
                    changed.underline = True
                    
                    if i < len(revised_text) - 1:
                        paragraph.add_run(revised_text[i+1:])
                    break
            return
        
        # Comparação detalhada palavra por palavra
        orig_words = original_text.split()
        rev_words = revised_text.split()
        
        s = difflib.SequenceMatcher(None, orig_words, rev_words)
        
        for tag, i1, i2, j1, j2 in s.get_opcodes():
            if tag == 'equal':
                # Palavras iguais - mantém normal
                paragraph.add_run(' '.join(orig_words[i1:i2]) + ' ')
            
            elif tag == 'replace':
                # Palavra mudou - mostra apenas a diferença
                old_word = ' '.join(orig_words[i1:i2])
                new_word = ' '.join(rev_words[j1:j2])
                
                # Compara caractere por caractere
                self._mark_word_difference(paragraph, old_word, new_word)
                paragraph.add_run(' ')
            
            elif tag == 'delete':
                # Removido
                removed = paragraph.add_run(' '.join(orig_words[i1:i2]))
                removed.font.strike = True
                removed.font.color.rgb = RGBColor(255, 0, 0)
                paragraph.add_run(' ')
            
            elif tag == 'insert':
                # Adicionado
                added = paragraph.add_run(' '.join(rev_words[j1:j2]))
                added.font.color.rgb = RGBColor(0, 128, 0)
                added.underline = True
                paragraph.add_run(' ')
    
    def _mark_word_difference(self, paragraph, old_word: str, new_word: str):
        """Marca diferença entre duas palavras mostrando só o que mudou"""
        
        # Encontra a diferença exata
        s = difflib.SequenceMatcher(None, old_word, new_word)
        
        for tag, i1, i2, j1, j2 in s.get_opcodes():
            if tag == 'equal':
                # Parte igual
                paragraph.add_run(old_word[i1:i2])
            
            elif tag == 'replace':
                # Parte diferente
                if i1 < i2:
                    old_part = paragraph.add_run(old_word[i1:i2])
                    old_part.font.strike = True
                    old_part.font.color.rgb = RGBColor(255, 0, 0)
                
                if j1 < j2:
                    new_part = paragraph.add_run(new_word[j1:j2])
                    new_part.font.color.rgb = RGBColor(0, 128, 0)
                    new_part.underline = True
            
            elif tag == 'delete':
                # Caracteres removidos
                removed = paragraph.add_run(old_word[i1:i2])
                removed.font.strike = True
                removed.font.color.rgb = RGBColor(255, 0, 0)
            
            elif tag == 'insert':
                # Caracteres adicionados
                added = paragraph.add_run(new_word[j1:j2])
                added.font.color.rgb = RGBColor(0, 128, 0)
                added.underline = True
    
    def _add_summary_with_dictionary(self, comparison_doc, all_corrections):
        """Adiciona sumário e dicionário completo no início do documento"""
        
        # Move para o início do documento
        first_para = comparison_doc.paragraphs[0] if comparison_doc.paragraphs else None
        
        # PARTE 1: TÍTULO E ESTATÍSTICAS
        title = comparison_doc.add_paragraph()
        title_run = title.add_run('RELATÓRIO COMPLETO DE REVISÃO')
        title_run.bold = True
        title_run.font.size = 16
        if first_para:
            title._element.addprevious(first_para._element)
        
        # Estatísticas
        total_changes = len(all_corrections)
        changes_by_type = {}
        for corr in all_corrections:
            corr_type = corr.get('type', 'outros')
            changes_by_type[corr_type] = changes_by_type.get(corr_type, 0) + 1
        
        stats = comparison_doc.add_paragraph()
        stats.add_run(f'\nTotal de correções: {total_changes}\n\n')
        stats.add_run('Por tipo:\n')
        for tipo, count in sorted(changes_by_type.items(), key=lambda x: x[1], reverse=True):
            stats.add_run(f'• {tipo.capitalize()}: {count}\n')
        if first_para:
            stats._element.addprevious(first_para._element)
        
        # Legenda
        legend = comparison_doc.add_paragraph()
        legend.add_run('\nLegenda: ')
        removed = legend.add_run('texto removido')
        removed.font.strike = True
        removed.font.color.rgb = RGBColor(255, 0, 0)
        legend.add_run(' | ')
        added = legend.add_run('texto adicionado')
        added.font.color.rgb = RGBColor(0, 128, 0)
        added.underline = True
        if first_para:
            legend._element.addprevious(first_para._element)
        
        # PARTE 2: DICIONÁRIO COMPLETO DE CORREÇÕES
        dict_break = comparison_doc.add_paragraph()
        dict_break.add_run('\n' + '='*80 + '\n')
        if first_para:
            dict_break._element.addprevious(first_para._element)
        
        dict_title = comparison_doc.add_paragraph()
        dict_title_run = dict_title.add_run('DICIONÁRIO DE TODAS AS CORREÇÕES')
        dict_title_run.bold = True
        dict_title_run.font.size = 14
        if first_para:
            dict_title._element.addprevious(first_para._element)
        
        # Organiza por página
        corrections_by_page = {}
        for corr in all_corrections:
            page = corr.get('page', 0)
            if page not in corrections_by_page:
                corrections_by_page[page] = []
            corrections_by_page[page].append(corr)
        
        # Lista correções página por página
        for page in sorted(corrections_by_page.keys()):
            if corrections_by_page[page]:
                # Cabeçalho da página
                page_header = comparison_doc.add_paragraph()
                page_run = page_header.add_run(f'\nPÁGINA {page}:')
                page_run.bold = True
                page_run.underline = True
                if first_para:
                    page_header._element.addprevious(first_para._element)
                
                # Correções da página
                for i, corr in enumerate(corrections_by_page[page], 1):
                    corr_para = comparison_doc.add_paragraph()
                    
                    # Localização
                    corr_para.add_run(f"{corr['location']}: ").bold = True
                    
                    # Tipo
                    corr_para.add_run(f"[{corr['type']}] ").font.color.rgb = RGBColor(0, 0, 139)
                    
                    # Erro → Correção
                    if corr['error'] not in ['[faltava ponto final]', '[faltava vírgula]', '[faltando]']:
                        error_run = corr_para.add_run(f'"{corr["error"]}"')
                        error_run.font.strike = True
                        error_run.font.color.rgb = RGBColor(255, 0, 0)
                    else:
                        corr_para.add_run(corr['error']).font.color.rgb = RGBColor(128, 128, 128)
                    
                    corr_para.add_run(' → ')
                    
                    fix_run = corr_para.add_run(f'"{corr["correction"]}"')
                    fix_run.font.color.rgb = RGBColor(0, 128, 0)
                    fix_run.underline = True
                    
                    if first_para:
                        corr_para._element.addprevious(first_para._element)
        
        # Linha final antes do documento
        final_line = comparison_doc.add_paragraph()
        final_line.add_run('\n' + '='*80)
        final_line.add_run('\nDOCUMENTO REVISADO COM MARCAÇÕES:\n\n')
        if first_para:
            final_line._element.addprevious(first_para._element)